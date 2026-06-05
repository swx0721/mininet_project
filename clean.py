from docx import Document
import re

input_file = "项目报告——去封皮版.docx"
output_file = "项目报告去除空格版.docx"

doc = Document(input_file)


def clean_text(text):

    # 1. 中文 <-> 英文/数字 去空格
    text = re.sub(r'([\u4e00-\u9fff])\s+([A-Za-z0-9%])', r'\1\2', text)
    text = re.sub(r'([A-Za-z0-9%])\s+([\u4e00-\u9fff])', r'\1\2', text)

    # 2. 数字内部空格（0 . 547 -> 0.547）
    text = re.sub(r'(\d)\s*\.\s*(\d)', r'\1.\2', text)

    # 3. 百分号前后空格（37.45 % -> 37.45%）
    text = re.sub(r'\s*%\s*', '%', text)

    # 4. mAP@0.5 : 0.95 -> mAP@0.5:0.95
    text = re.sub(r'\s*:\s*', ':', text)

    # 5. 中文标点前后空格
    text = re.sub(r'\s*([，。；：！？、】【（）])\s*', r'\1', text)

    # 6. 多余空格压缩
    text = re.sub(r'\s{2,}', ' ', text)

    return text


# ===== 清洗正文 =====
for para in doc.paragraphs:
    if para.text.strip():
        para.text = clean_text(para.text)

# ===== 清洗表格 =====
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.text = clean_text(para.text)

doc.save(output_file)

print("清洗完成：", output_file)